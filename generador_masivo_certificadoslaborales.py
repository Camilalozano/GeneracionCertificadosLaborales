# ============================================================
# GENERADOR AUTOMÁTICO DE CERTIFICADOS LABORALES ATENEA
# Input: Excel .xlsx con cualquier nombre y cualquier nombre de hoja. Se toma siempre la primera pestaña.
# Output: ZIP con certificados .docx y, si es posible, .pdf
# ============================================================

import os
import re
import zipfile
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# FUNCIONES AUXILIARES
# =========================

def limpiar_nombre_archivo(texto):
    texto = str(texto).strip()
    texto = re.sub(r'[\\/*?:"<>|]', "_", texto)
    texto = re.sub(r"\s+", "_", texto)
    return texto[:180]


def valor_limpio(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def valor_documento_texto(valor):
    """
    Normaliza números de documento leídos desde Excel para usarlos como texto.

    Cuando pandas interpreta la columna como numérica, documentos como
    1032373832 pueden llegar como 1032373832.0. Para evitar que ese sufijo
    aparezca en el certificado o en el nombre del archivo, los flotantes que
    representan enteros se convierten explícitamente a texto entero.
    """
    if pd.isna(valor):
        return ""

    if isinstance(valor, float) and valor.is_integer():
        return str(int(valor))

    texto = str(valor).strip()
    if re.fullmatch(r"\d+\.0+", texto):
        return texto.split(".", 1)[0]

    return texto




def limpiar_documento_para_archivo(valor):
    """
    Limpia el documento para usarlo en el nombre del archivo.
    Ejemplo: 1.032.373.832 -> 1032373832.
    También corrige valores numéricos leídos desde Excel como 1032373832.0.
    """
    texto = valor_documento_texto(valor)
    if not texto:
        return ""

    # Si el documento tiene puntos, comas, espacios o guiones, conserva solo dígitos.
    solo_digitos = re.sub(r"\D", "", texto)
    return solo_digitos if solo_digitos else limpiar_nombre_archivo(texto)

def formatear_documento_colombiano(valor):
    """
    Formatea documentos para que se vean como en la plantilla:
    1032373832 -> 1.032.373.832.
    Si el valor contiene letras o caracteres especiales, conserva el texto limpio.
    """
    texto = valor_documento_texto(valor)
    if not texto:
        return ""

    solo_digitos = re.sub(r"\D", "", texto)
    if solo_digitos and len(solo_digitos) == len(re.sub(r"\s", "", texto)):
        return f"{int(solo_digitos):,}".replace(",", ".")

    return texto


def formatear_valor_pesos(valor):
    """
    Formatea valores monetarios provenientes de SECOP para el certificado.
    Ejemplo: 90000000 -> $90.000.000 M/CTE.
    """
    if pd.isna(valor) or str(valor).strip() == "":
        return ""

    texto = str(valor).strip()

    # Si ya viene con símbolo o texto, limpia espacios y conserva el contenido.
    if "$" in texto or "M/CTE" in texto.upper() or "PESOS" in texto.upper():
        return re.sub(r"\s+", " ", texto)

    try:
        numero = pd.to_numeric(texto.replace(".", "").replace(",", "."), errors="coerce")
        if pd.isna(numero):
            return texto
        numero_entero = int(round(float(numero)))
        return f"${numero_entero:,}".replace(",", ".") + " M/CTE"
    except Exception:
        return texto


def formatear_fecha_larga(valor):
    """
    Formatea fechas en estilo de certificado: 17 de marzo del 2025.
    """
    if pd.isna(valor) or str(valor).strip() == "":
        return ""

    texto = str(valor).split("((")[0].strip()
    fecha = pd.to_datetime(texto, dayfirst=True, errors="coerce")
    if pd.isna(fecha):
        return texto

    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    return f"{fecha.day} de {meses[fecha.month]} del {fecha.year}"


def formatear_fecha(valor):
    if pd.isna(valor) or str(valor).strip() == "":
        return ""

    texto = str(valor)

    # Limpia textos tipo SECOP con zona horaria
    texto = texto.split("((")[0].strip()

    try:
        fecha = pd.to_datetime(texto, dayfirst=True, errors="coerce")
        if pd.isna(fecha):
            return texto
        return fecha.strftime("%d/%m/%Y")
    except Exception:
        return texto


def definir_calidad(tipo_contrato):
    tipo = valor_limpio(tipo_contrato).upper()

    if "PRESTACION DE SERVICIOS PROFESIONALES" in tipo or "PRESTACIÓN DE SERVICIOS PROFESIONALES" in tipo:
        return "en calidad de contratista"
    elif "PRESTACIÓN DE SERVICIOS" in tipo or "PRESTACION DE SERVICIOS" in tipo:
        return "en calidad de contratista"
    else:
        return "en calidad de contratista"


def agregar_parrafo(doc, texto="", bold=False, align=None, size=11):
    p = doc.add_paragraph()
    if align:
        p.alignment = align

    run = p.add_run(texto)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    return p


def agregar_titulo_centrado(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def agregar_campo(doc, etiqueta, valor):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)

    r1 = p.add_run(f"{etiqueta}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(11)

    r2 = p.add_run(valor)
    r2.font.name = "Arial"
    r2.font.size = Pt(11)


def agregar_hipervinculo(parrafo, texto, url, color="FF0000", bold=True):
    relacion_id = parrafo.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relacion_id)

    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    if bold:
        run_properties.append(OxmlElement("w:b"))

    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_properties.append(run_color)

    run_font = OxmlElement("w:rFonts")
    run_font.set(qn("w:ascii"), "Arial")
    run_font.set(qn("w:hAnsi"), "Arial")
    run_properties.append(run_font)

    run_size = OxmlElement("w:sz")
    run_size.set(qn("w:val"), "22")
    run_properties.append(run_size)

    new_run.append(run_properties)

    run_text = OxmlElement("w:t")
    run_text.text = texto
    new_run.append(run_text)

    hyperlink.append(new_run)
    parrafo._p.append(hyperlink)


def agregar_aviso_automatizado(doc, url):
    texto_aviso = (
        "Esta información ha sido generada mediante el uso de modelos de machine learning "
        "y procesos automatizados de extracción de datos. Se recomienda validar y verificar "
        "su contenido con las fuentes oficiales disponibles en el siguiente enlace antes de su uso."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)

    run_aviso = p.add_run(texto_aviso)
    run_aviso.bold = True
    run_aviso.font.name = "Arial"
    run_aviso.font.size = Pt(11)
    run_aviso.font.color.rgb = RGBColor(255, 0, 0)

    url = valor_limpio(url)
    if url:
        p.add_run(" ")
        agregar_hipervinculo(p, url, url)


def normalizar_obligaciones(texto):
    texto = valor_limpio(texto)

    if not texto:
        return []

    # Divide por numeración tipo 1. 2. 3.
    partes = re.split(r"\n?\s*(?=\d+\.\s)", texto)
    partes = [p.strip() for p in partes if p.strip()]

    if len(partes) <= 1:
        # Si no viene numerado, divide por saltos de línea
        partes = [p.strip() for p in texto.split("\n") if p.strip()]

    return partes


def configurar_margenes(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def agregar_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(
        "Cr 10 # 28-49. Torre A, piso 26.\n"
        "Bogotá D.C. Colombia\n"
        "(601) 666 0006\n"
        "www.agenciaatenea.gov.co"
    )
    run.font.name = "Arial"
    run.font.size = Pt(8)


def convertir_docx_a_pdf(docx_path, output_dir):
    """
    Intenta convertir DOCX a PDF usando LibreOffice.
    Si no está instalado, simplemente omite la conversión.
    """
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")

    if not libreoffice:
        return None

    try:
        subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path)
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )

        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        return pdf_path if pdf_path.exists() else None

    except Exception:
        return None


# =========================
# FUNCIÓN PRINCIPAL
# =========================

def crear_certificado(row, output_dir, logo_path=None):
    # Variables tomadas de la base SECOP_NoEstructurado / contratos electrónicos
    numero_contrato = valor_limpio(row.get("referencia_del_contrato (contratos_electronicos)", ""))
    entidad_adjudicataria = valor_limpio(row.get("proveedor_adjudicado (contratos_electronicos)", "")).upper()
    documento = formatear_documento_colombiano(row.get("documento_proveedor (contratos_electronicos)", ""))
    objeto = valor_limpio(row.get("descripcion_del_proceso (contratos_electronicos)", ""))
    obligaciones = valor_limpio(row.get("obligaciones específicas consolidadas", ""))
    valor_total = formatear_valor_pesos(row.get("valor_del_contrato (contratos_electronicos)", ""))
    fecha_inicio = formatear_fecha_larga(row.get("fecha_de_inicio_del_contrato (contratos_electronicos)", ""))
    fecha_fin = formatear_fecha_larga(row.get("fecha_de_fin_del_contrato (contratos_electronicos)", ""))
    tiempo_ejecucion = valor_limpio(row.get("duracion_contrato", ""))
    tipo_contrato = valor_limpio(row.get("justificacion_modalidad_de (contratos_electronicos)", ""))
    estado = valor_limpio(row.get("estado_contrato (contratos_electronicos)", ""))
    url = valor_limpio(row.get("urlproceso (contratos_electronicos)", ""))
    calidad = definir_calidad(row.get("justificacion_modalidad_de (contratos_electronicos)", ""))

    doc = Document()
    configurar_margenes(doc)

    # Logo superior izquierdo
    if logo_path and Path(logo_path).exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(logo_path), width=Inches(1.45))

    agregar_titulo_centrado(
        doc,
        "LA SUBGERENCIA DE GESTIÓN ADMINISTRATIVA DE LA AGENCIA DISTRITAL PARA LA\n"
        "EDUCACIÓN SUPERIOR, LA CIENCIA Y LA TECNOLOGÍA - ATENEA"
    )

    agregar_titulo_centrado(doc, "CERTIFICA QUE:")

    texto_intro = (
        f"Revisados los archivos de contratación, se encontró que {entidad_adjudicataria}, "
        f"quien se identifica con cédula de ciudadanía {documento}, "
        f"suscribió con la Agencia Distrital para la Educación Superior, "
        f"la Ciencia y la Tecnología, {calidad}, el siguiente contrato:"
    )

    p = agregar_parrafo(doc, texto_intro, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    agregar_titulo_centrado(doc, f"CONTRATO No. {numero_contrato}")

    agregar_campo(doc, "Objeto", objeto)
    agregar_campo(doc, "Valor total", valor_total)
    agregar_campo(doc, "Fecha de inicio", fecha_inicio)
    agregar_campo(doc, "Fecha de finalización", fecha_fin)
    agregar_campo(doc, "Tiempo de ejecución", tiempo_ejecucion)
    agregar_campo(doc, "Estado", estado)

    agregar_parrafo(doc, "Obligaciones específicas del contratista:", bold=True)

    lista_obligaciones = normalizar_obligaciones(obligaciones)

    for i, obligacion in enumerate(lista_obligaciones, start=1):
        obligacion = re.sub(r"^\d+\.\s*", "", obligacion).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

        r = p.add_run(f"{i}. {obligacion}")
        r.font.name = "Arial"
        r.font.size = Pt(10.5)

    fecha_hoy = datetime.now().strftime("%d/%m/%Y")
    agregar_parrafo(doc, f"La solicitud se expide en Bogotá, el día {fecha_hoy}.")
    agregar_parrafo(doc, "Cordialmente,")

    doc.add_paragraph("\n")

    agregar_parrafo(doc, "CAMILO CARDOZO CRUZ", bold=True)
    agregar_parrafo(doc, "Subgerente de Gestión Administrativa")

    agregar_parrafo(doc, "Proyectó. Revisó.", size=8)

    agregar_aviso_automatizado(doc, url)

    agregar_footer(doc)

    documento_archivo = limpiar_documento_para_archivo(row.get("numero_documento_contratista", ""))
    nombre_archivo = limpiar_nombre_archivo(f"{numero_contrato}_{documento_archivo}")
    docx_path = output_dir / f"{nombre_archivo}.docx"
    doc.save(docx_path)

    return docx_path


def main():
    print("=== GENERADOR DE CERTIFICADOS LABORALES ATENEA ===")

    ruta_excel = input("Ingresa la ruta completa del archivo Excel .xlsx de entrada: ").strip().strip('"')
    ruta_salida = input("Ingresa la ruta de la carpeta donde deseas guardar el ZIP: ").strip().strip('"')
    ruta_logo = input("Ingresa la ruta del logo de Atenea en PNG/JPG. Si no tienes logo, deja vacío y presiona Enter: ").strip().strip('"')

    ruta_excel = Path(ruta_excel)
    ruta_salida = Path(ruta_salida)
    ruta_salida.mkdir(parents=True, exist_ok=True)

    if not ruta_excel.exists():
        raise FileNotFoundError(f"No se encontró el archivo Excel: {ruta_excel}")

    print("\nLeyendo base de datos...")

    # El archivo de entrada puede tener cualquier nombre y la hoja puede tener
    # cualquier nombre. Por instrucción, se toma siempre la primera pestaña.
    df = pd.read_excel(
        ruta_excel,
        sheet_name=0,
        dtype={
            "documento_proveedor (contratos_electronicos)": "string",
            "numero_documento_contratista": "string",
        },
    )

    columna_obligaciones = "obligaciones específicas consolidadas"

    columnas_requeridas = [
        "referencia_del_contrato (contratos_electronicos)",
        "proveedor_adjudicado (contratos_electronicos)",
        "documento_proveedor (contratos_electronicos)",
        "descripcion_del_proceso (contratos_electronicos)",
        columna_obligaciones,
        "valor_del_contrato (contratos_electronicos)",
        "fecha_de_inicio_del_contrato (contratos_electronicos)",
        "fecha_de_fin_del_contrato (contratos_electronicos)",
        "justificacion_modalidad_de (contratos_electronicos)",
        "estado_contrato (contratos_electronicos)",
        "urlproceso (contratos_electronicos)",
        "numero_documento_contratista",
    ]

    for columna_requerida in columnas_requeridas:
        if columna_requerida not in df.columns:
            raise ValueError(f"No existe la columna requerida: {columna_requerida}")

    df_filtrado = df[
        df[columna_obligaciones].notna()
        & (df[columna_obligaciones].astype(str).str.strip() != "")
    ].copy()

    print(f"Registros con obligaciones no faltantes: {len(df_filtrado)}")

    carpeta_certificados = ruta_salida / "certificados_laborales_atenea"
    carpeta_certificados.mkdir(parents=True, exist_ok=True)

    certificados_generados = []

    for idx, row in df_filtrado.iterrows():
        try:
            docx_path = crear_certificado(
                row=row,
                output_dir=carpeta_certificados,
                logo_path=ruta_logo if ruta_logo else None
            )

            certificados_generados.append(docx_path)

            pdf_path = convertir_docx_a_pdf(docx_path, carpeta_certificados)
            if pdf_path:
                certificados_generados.append(pdf_path)

            print(f"Certificado generado: {docx_path.name}")

        except Exception as e:
            print(f"Error en registro {idx}: {e}")

    zip_path = ruta_salida / "certificados_laborales_atenea.zip"

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for archivo in certificados_generados:
            zipf.write(archivo, arcname=archivo.name)

    print("\n===================================")
    print("Proceso finalizado correctamente.")
    print(f"ZIP generado en: {zip_path}")
    print("===================================")


if __name__ == "__main__":
    main()
